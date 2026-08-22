# -*- coding: utf-8 -*-
"""Dựng tài liệu Bản đồ database Selliotech (.docx).

Mọi con số trong tài liệu đọc thẳng từ MySQL máy cục bộ và máy thật lúc chạy
lệnh, không chép lại từ trí nhớ.
"""

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

OUT = r"c:\huy\sass\sass-docx\selliotech-ban-do-database.docx"

XANH = RGBColor(0x1F, 0x3A, 0x93)
XAM = RGBColor(0x55, 0x5D, 0x6B)
DO = RGBColor(0xB3, 0x26, 0x1E)

doc = Document()

# ---------------------------------------------------------------- kiểu chữ
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)


def p(text="", *, dam=False, mau=None, co=None, sau=6, truoc=0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(sau)
    par.paragraph_format.space_before = Pt(truoc)
    if text:
        r = par.add_run(text)
        r.bold = dam
        if mau is not None:
            r.font.color.rgb = mau
        if co is not None:
            r.font.size = Pt(co)
    return par


def rich(parts, *, sau=6, truoc=0, style=None):
    """parts: list các tuple (chữ, đậm, mã-nguồn)."""
    par = doc.add_paragraph(style=style)
    par.paragraph_format.space_after = Pt(sau)
    par.paragraph_format.space_before = Pt(truoc)
    for text, dam, ma in parts:
        r = par.add_run(text)
        r.bold = dam
        if ma:
            r.font.name = "Consolas"
            r.font.size = Pt(10)
    return par


def cham(parts):
    return rich(parts, style="List Bullet", sau=3)


def bang(tieu_de, dong, rong=None):
    t = doc.add_table(rows=1, cols=len(tieu_de))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, ten in enumerate(tieu_de):
        o = t.rows[0].cells[i]
        o.text = ""
        r = o.paragraphs[0].add_run(ten)
        r.bold = True
        r.font.size = Pt(10)
    for d in dong:
        cells = t.add_row().cells
        for i, gt in enumerate(d):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            # gt dạng "`code`" -> chữ mã nguồn
            for khuc, la_ma in tach_ma(str(gt)):
                r = par.add_run(khuc)
                r.font.size = Pt(10)
                if la_ma:
                    r.font.name = "Consolas"
                    r.font.size = Pt(9.5)
    if rong:
        for hang in t.rows:
            for i, w in enumerate(rong):
                hang.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def tach_ma(s):
    """Cắt chuỗi theo dấu ` thành các khúc (chữ, có phải mã nguồn không)."""
    ra, dem = [], s.split("`")
    for i, khuc in enumerate(dem):
        if khuc:
            ra.append((khuc, i % 2 == 1))
    return ra or [(s, False)]


def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = XANH
    return h


def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = XANH
    return h


def khung(nhan, noi_dung, mau=DO):
    """Một dòng cảnh báo/ghi chú nổi bật."""
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(8)
    r = par.add_run(nhan + " ")
    r.bold = True
    r.font.color.rgb = mau
    for khuc, la_ma in tach_ma(noi_dung):
        rr = par.add_run(khuc)
        if la_ma:
            rr.font.name = "Consolas"
            rr.font.size = Pt(10)
    return par


# ---------------------------------------------------------------- trang đầu
tieu = doc.add_heading("Selliotech — Bản đồ database", level=0)
for r in tieu.runs:
    r.font.color.rgb = XANH
p("Từng bảng có ý nghĩa gì, làm chức năng gì — trên máy cục bộ và trên máy thật.",
  co=12, mau=XAM)

bang(
    ["", ""],
    [
        ["Ngày viết", "12/08/2026"],
        ["Máy cục bộ", "`127.0.0.1:3306` — MySQL của XAMPP"],
        ["Máy thật", "`103.78.2.230:3306` — VPS, bản đang chạy là commit `5e3b946`"],
        ["Nguồn số liệu", "Đọc thẳng từ MySQL của cả hai máy lúc viết tài liệu này, không chép lại từ tài liệu cũ"],
        ["Số bảng", "`selliotech`: 43 bảng (giống hệt nhau ở hai máy) · `selliotech_platform`: 7 bảng ở máy cục bộ, 5 bảng ở máy thật"],
    ],
    rong=[1.3, 5.2],
)

khung("Đọc mục 5 trước khi sửa gì:",
      "hai máy đang KHÁC nhau — `apps` và `plans` mới có ở máy cục bộ, chưa deploy.")

doc.add_page_break()

# ---------------------------------------------------------------- 1
h1("1. Đọc trước: hai database, hai vòng đời")

p("Hệ thống không có một database duy nhất. Nó có hai, và tách hẳn nhau là chủ ý chứ không phải lịch sử để lại:")

bang(
    ["Database", "Tên gọi", "Chứa gì"],
    [
        ["`selliotech`", "Data plane — mặt bán hàng",
         "Dữ liệu KINH DOANH CỦA KHÁCH: sản phẩm, đơn hàng, tồn kho, nhân viên, khách mua. 43 bảng."],
        ["`selliotech_platform`", "Control plane — sổ cái nền tảng",
         "Dữ liệu VỀ KHÁCH: ai đã mua phần mềm, mua gói nào, còn hạn tới bao giờ, tên miền nào, và mình bán những phần mềm gì."],
    ],
    rong=[1.6, 1.7, 3.2],
)

p("Vì sao không nhét chung vào một database:", dam=True, truoc=4)
cham([("Hai vòng đời khác nhau. ", True, False),
      ("Dữ liệu bán hàng của một khách có thể phải tách ra máy chủ riêng, khôi phục riêng, hoặc xoá đi khi khách nghỉ. Sổ nền tảng thì không bao giờ đi theo một khách nào.", False, False)])
cham([("Ranh giới quyền. ", True, False),
      ("Tài khoản MySQL của khu bán hàng không cần — và không nên — đọc nổi bảng tiền nong của toàn bộ khách hàng khác.", False, False)])
cham([("Đường đi tiếp. ", True, False),
      ("Khi một khách lớn cần database riêng, data plane nhân bản ra nhiều database, còn control plane vẫn đúng một.", False, False)])

khung("Không có khoá ngoại nào bắc qua hai database.",
      "MySQL cho phép, nhưng ở đây tuyệt đối không dùng: nó biến hai database thành một khối dính liền, "
      "không tách máy chủ được, không khôi phục riêng được, và `mysqldump` từng cái nạp lại sẽ không lên vì thiếu bảng cha. "
      "Ghép dữ liệu hai bên là việc của tầng Go.", mau=DO)

khung("Cạm bẫy hay gặp nhất khi mở DBeaver:", "bảng `tenants` có ở CẢ HAI database và là hai bảng khác nhau. "
      "Gõ nhầm database thì câu truy vấn vẫn chạy, chỉ là ra dữ liệu của bảng khác. Xem mục 3.3.", mau=DO)

# ---------------------------------------------------------------- 2
h1("2. Ở đâu có gì")

p("Danh sách database nhìn thấy trong DBeaver, và cái nào là thật:")

bang(
    ["Database", "Máy cục bộ", "Máy thật", "Là gì"],
    [
        ["`selliotech`", "Có", "Có", "Dữ liệu bán hàng. Database chính của phần mềm."],
        ["`selliotech_platform`", "Có", "Có", "Sổ cái nền tảng. Database chính thứ hai."],
        ["`selliotech_tenant_test`", "Có", "Không", "Rác của bộ kiểm ranh giới hai khách hàng — bộ test chạy trên API thật nên cần database riêng để ghi đè thoải mái."],
        ["`selliotech_tenant_test_nen_tang`", "Có", "Không", "Control plane của chính bộ test đó. Tên suy ra từ tên trên + hậu tố `_nen_tang`."],
        ["`test`", "Có", "Không", "Database rỗng XAMPP tạo sẵn lúc cài. Không liên quan tới dự án."],
    ],
    rong=[1.9, 0.8, 0.8, 3.0],
)

khung("Hai database test KHÔNG được xuất hiện trên máy thật.", "Thấy chúng ở đó nghĩa là có người vừa chạy test lên máy thật.", mau=DO)

doc.add_page_break()

# ---------------------------------------------------------------- 3
h1("3. selliotech_platform — sổ cái nền tảng")

p("Bảy bảng. Đây là phần trả lời câu hỏi kinh doanh: mình bán gì, giá bao nhiêu, ai đang dùng, còn hạn tới bao giờ.")

bang(
    ["Bảng", "Trả lời câu hỏi", "Máy cục bộ", "Máy thật"],
    [
        ["`apps`", "Mình bán những phần mềm nào?", "1 dòng", "CHƯA CÓ BẢNG"],
        ["`plans`", "Mỗi phần mềm có những gói giá nào?", "3 dòng", "CHƯA CÓ BẢNG"],
        ["`tenants`", "Khách hàng của mình là ai?", "0 dòng", "0 dòng"],
        ["`subscriptions`", "Khách đang dùng gói nào, tới bao giờ?", "0 dòng", "0 dòng"],
        ["`tenant_domains`", "Tên miền nào thuộc cửa hàng nào?", "0 dòng", "0 dòng"],
        ["`platform_users`", "Ai được vào khu điều hành?", "0 dòng", "0 dòng"],
        ["`schema_migrations`", "Lược đồ này đã chạy tới đâu?", "4 dòng", "1 dòng"],
    ],
    rong=[1.5, 2.9, 1.0, 1.1],
)

# ---- 3.1 apps
h2("3.1  apps — danh mục phần mềm mình bán")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Mỗi dòng là MỘT phần mềm nền tảng bán ra. Đây là câu trả lời cho “bán cái gì”."],
        ["Chức năng", "Làm bảng cha của `plans` (bảng giá gắn theo từng app), và về sau là chỗ để biết một khách đang mua phần mềm nào khi có sản phẩm thứ hai."],
        ["Cột đáng chú ý", "`code` (tên ngắn chữ thường, trùng tiền tố tên miền: `order` → `order.selliotech.store`) · `name` · `tagline` · `status` = `planned` | `active` | `retired`"],
        ["Máy cục bộ", "1 dòng: `order` — “Sellio Order” — `active` (chính phần mềm quản trị bán hàng đang chạy)"],
        ["Máy thật", "Chưa có bảng — migration `0002` chưa deploy"],
    ],
    rong=[1.3, 5.2],
)
cham([("`code` chứ không phải `id` là thứ đem đi dùng ở nơi khác (cấu hình, tên miền, về sau là JWT/URL): `id` là số tự sinh của MỘT database, chép cấu hình sang máy khác là lệch.", False, False)])
cham([("`status` mặc định là `planned` chứ không phải `active`: dòng mới thêm mà tự nhiên bán được là cách nhanh nhất để ký hợp đồng cho một phần mềm chưa tồn tại.", False, False)])
cham([("`retired` nghĩa là NGỪNG BÁN MỚI, không phải ngừng chạy — khách cũ vẫn dùng tiếp.", False, False)])

# ---- 3.2 plans
h2("3.2  plans — bảng giá hiện hành của từng app")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Mỗi dòng là MỘT MỨC GIÁ: gói này, của app này, theo chu kỳ này."],
        ["Chức năng", "Giữ giá niêm yết và hạn mức của từng gói, để khu điều hành và (sau này) trang landing đọc từ một chỗ duy nhất thay vì gõ tay ở hai nơi."],
        ["Cột đáng chú ý", "`app_id` → `apps.id` · `code` (mã gói) · `name` · `billing_cycle` = `thang` | `nam` · `price` · `max_shops` · `own_domain` · `trial_days` · `status` = `active` | `retired`"],
        ["Khoá duy nhất", "(`app_id`, `code`, `billing_cycle`) — một gói bán cả tháng lẫn năm là HAI dòng, hai giá. Chu kỳ là dữ liệu, không phải tên cột."],
        ["Máy cục bộ", "3 dòng cho app `order`, chép đúng bảng giá đang công khai trên landing"],
        ["Máy thật", "Chưa có bảng — migration `0003` chưa deploy"],
    ],
    rong=[1.3, 5.2],
)
p("Ba dòng đang có:", dam=True, truoc=4)
bang(
    ["Mã gói", "Tên", "Chu kỳ", "Giá (VND)", "max_shops", "Tên miền riêng"],
    [
        ["`khoi_dau`", "Khởi đầu", "thang", "199.000", "1", "0 — không"],
        ["`cua_hang`", "Cửa hàng", "thang", "499.000", "1", "0 — không"],
        ["`chuoi`", "Chuỗi", "thang", "NULL = Liên hệ", "NULL = chốt lúc ký", "1 — CÓ"],
    ],
    rong=[1.0, 0.9, 0.7, 1.3, 1.4, 1.0],
)
p("Cả ba gói đều có 14 ngày dùng thử (`trial_days`).", sau=8)
khung("Đây là BẢNG GIÁ, không phải HỢP ĐỒNG.",
      "`subscriptions.price` và `subscriptions.max_shops` là bản CHÉP lúc ký và từ đó sống độc lập. "
      "Vì vậy tuyệt đối không thêm `subscriptions.plan_id` trỏ sang đây: khoá ngoại sẽ biến hợp đồng đã ký "
      "thành thứ đi theo bảng giá hiện hành, tức là đổi giá tháng sau là tăng giá luôn người đang trả tiền.", mau=DO)
cham([("`price` NULL nghĩa là “Liên hệ”, KHÁC 0 (miễn phí). Bên `subscriptions.price` thì NOT NULL: không ai ký hợp đồng “liên hệ”.", False, False)])
cham([("Sửa giá là `UPDATE` thẳng dòng, không cần migration mới — bảng này giữ giá HIỆN HÀNH, lịch sử giá nằm ở `subscriptions.price` của từng hợp đồng.", False, False)])
cham([("Sửa tay bằng SQL thì nhớ set `updated_at = NOW(3)`: cột đó không có `ON UPDATE CURRENT_TIMESTAMP`, cả repo giao cho GORM điền.", False, False)])
cham([("Thêm mã gói mới thì phải mở rộng ENUM `subscriptions.plan` trong CÙNG một migration, nếu không sẽ bán được gói mà MySQL từ chối ghi thuê bao.", False, False)])
cham([("Giá còn được gõ tay lần thứ hai trong `landing_shop/index.html`. Đổi giá là phải sửa cả hai chỗ, tới ngày landing đọc qua API mới hết.", False, False)])
cham([("`own_domain` là ĐIỀU KHOẢN BÁN HÀNG nằm trong dữ liệu, không phải luật kỹ thuật: đổi chính sách là `UPDATE` một ô, không phải sửa code rồi triển khai lại. Chi tiết ở mục 3.5.", False, False)])

# ---- 3.3 tenants
h2("3.3  tenants — sổ đăng ký khách hàng")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Mỗi dòng là một khách hàng đã mua phần mềm, nhìn từ phía kinh doanh: họ là ai, liên lạc với ai, đang mở hay đang khoá."],
        ["Chức năng", "Chỗ RA QUYẾT ĐỊNH mở/khoá một cửa hàng, và là bảng cha của `subscriptions` với `tenant_domains`."],
        ["Cột đáng chú ý", "`id` (ghi tay!) · `code` (mã cửa hàng khách gõ lúc đăng nhập) · `name` · `status` = `active` | `suspended` · `contact_*` · `note`"],
        ["Máy cục bộ", "0 dòng"],
        ["Máy thật", "0 dòng — trong khi `selliotech.tenants` đã có 2 cửa hàng thật (`cuahang`, `quochuy`)"],
    ],
    rong=[1.3, 5.2],
)
khung("Bảng `tenants` tồn tại ở CẢ HAI database và không phải bản sao thừa:", "", mau=DO)
cham([("`selliotech.tenants` — ", False, False), ("bảng CHA", True, False),
      (" của mọi khoá ngoại `tenant_id` trong dữ liệu bán hàng, và là thứ API đọc lúc đăng nhập 3 ô. Bỏ nó đi thì 37 bảng mất khoá ngoại.", False, False)])
cham([("`selliotech_platform.tenants` — ", False, False), ("sổ đăng ký", True, False),
      (" của nền tảng: khách này là ai, gắn với gói nào, tên miền nào.", False, False)])
cham([("`status` có ở cả hai và PHẢI ghi cả hai lượt: control plane là nơi ra quyết định, data plane là nơi API thực sự đọc lúc chặn đăng nhập. Đổi một bên mà quên bên kia thì khách đã ngừng trả tiền vẫn vào bán hàng bình thường.", False, False)])
cham([("`id` ở bảng này KHÔNG tự tăng — phải ghi thẳng đúng id đã có bên `selliotech.tenants`, vì hôm nay việc cấp id vẫn do data plane làm. Để nó tự sinh thì hai bên đếm độc lập và tenant 7 bên này là tenant 9 bên kia.", False, False)])

# ---- 3.4 subscriptions
h2("3.4  subscriptions — khách đang dùng gói nào, tới bao giờ")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Hợp đồng đang chạy của một khách: gói gì, giá bao nhiêu, chu kỳ nào, hết hạn ngày nào."],
        ["Chức năng", "Nguồn để biết ai còn hạn ai hết hạn — câu hỏi chạy hằng ngày là “gói nào sắp hết hạn hoặc đã quá hạn mà chưa xử lý”."],
        ["Cột đáng chú ý", "`tenant_id` · `plan` = `khoi_dau` | `cua_hang` | `chuoi` · `status` = `trial` | `active` | `past_due` | `canceled` · `billing_cycle` · `price` · `max_shops` · `started_at` / `ends_at` / `trial_ends_at` / `canceled_at`"],
        ["Máy cục bộ", "0 dòng"],
        ["Máy thật", "0 dòng"],
    ],
    rong=[1.3, 5.2],
)
cham([("Một tenant chỉ có TỐI ĐA MỘT thuê bao còn hiệu lực, và điều đó do khoá dưới database ép chứ không do tầng Go tự nhớ: cột `current_mark` bằng 1 khi còn hiệu lực và NULL khi đã huỷ, nên lịch sử gói cũ nằm lại thoải mái.", False, False)])
cham([("Gia hạn = đẩy `ends_at` của chính dòng đó. Đổi gói = huỷ dòng cũ (`status='canceled'` + `canceled_at`) rồi thêm dòng mới.", False, False)])
cham([("`price` và `max_shops` là bản chốt với khách lúc ký — đọc thẳng ở dòng này, đừng suy ra từ `plan` hay tra sang `plans`.", False, False)])
cham([("Còn thiếu `app_id`: hôm nay mọi thuê bao ngầm hiểu là thuê bao của app `order`. Nên bù sớm, vì bảng đang 0 dòng nên điền ngược không tốn gì.", False, False)])

# ---- 3.5 tenant_domains
h2("3.5  tenant_domains — tên miền nào thuộc cửa hàng nào")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Ánh xạ một tên miền sang một cửa hàng: vào `cuahangabc.selliotech.store` thì biết ngay đang đứng ở tiệm nào, không phải gõ mã cửa hàng."],
        ["Chức năng", "Cửa duy nhất mà luồng phục vụ khách VÃNG LAI đọc sang control plane. Quản lý bằng lệnh `go run ./cmd/ten-mien them|xoa`."],
        ["Cột đáng chú ý", "`tenant_id` · `host` (chữ thường, không scheme, không cổng) · `kind` = `subdomain` | `custom` · `is_primary` · `verified_at` · `note` (lý do khi cấp đặc cách)"],
        ["Máy cục bộ", "0 dòng"],
        ["Máy thật", "0 dòng — hôm nay cả nền tảng vẫn chạy trên vài tên miền cố định và mã cửa hàng gõ tay lúc đăng nhập"],
    ],
    rong=[1.3, 5.2],
)
cham([("`host` UNIQUE toàn bảng: một tên miền không thể trỏ vào hai cửa hàng — đây chính là chỗ chặn khách này khai tên miền của khách kia.", False, False)])
cham([("Mỗi tenant có tối đa MỘT tên miền chính, ép bằng cột `primary_mark` chứ không bằng lời dặn.", False, False)])
cham([("`verified_at` NULL nghĩa là chưa xác minh DNS và CHƯA được xin chứng chỉ HTTPS: xin hỏng nhiều lần liên tiếp là Let's Encrypt khoá hạn mức cả tuần.", False, False)])
cham([("Bảng này mà đọc không được thì không tên miền nào phân giải được — cả trang bán hàng đứng im trong khi process vẫn báo khoẻ.", False, False)])

p("Ai được cấp tên miền riêng — chính sách đã chốt", dam=True, truoc=8)
p("Tên miền riêng là TÍNH NĂNG CỦA GÓI, không phải thứ mọi khách đều có:")
bang(
    ["Khách", "Địa chỉ dùng", "Đăng nhập"],
    [
        ["Đang dùng thử (mọi gói)", "`order.selliotech.store` dùng chung", "gõ mã cửa hàng ở ô 1"],
        ["Gói Khởi đầu / Cửa hàng đã trả tiền", "`order.selliotech.store` dùng chung", "gõ mã cửa hàng ở ô 1"],
        ["Gói Chuỗi đã trả tiền", "tên miền riêng: `quochuy.selliotech.store` hoặc tên miền của khách", "vào đúng địa chỉ là biết tiệm nào"],
    ],
    rong=[1.9, 2.7, 1.9],
)
p("Luật được ép ở `cmd/ten-mien them` — đường DUY NHẤT ghi vào sổ này. Ba điều kiện, thiếu một là từ chối:", truoc=4)
cham([("có thuê bao còn hiệu lực trong sổ nền tảng;", False, False)])
cham([("thuê bao đó `status = 'active'` — đã trả tiền. `trial` không được (đó chính là câu “dùng thử thì dùng chung”), `past_due` cũng không: đang nợ thì không cấp THÊM địa chỉ mới, còn địa chỉ đã cấp vẫn chạy;", False, False)])
cham([("gói của thuê bao đó có `plans.own_domain = 1`.", False, False)])
p("Địa chỉ được sinh TỰ ĐỘNG từ tên cửa hàng", dam=True, truoc=8)
p("Người vận hành không nghĩ ra địa chỉ và không gõ tay tên khách — chạy một lệnh:")
rich([("go run ./cmd/ten-mien tu-dong --ma-cua-hang quochuy", False, True)], sau=2)
rich([("go run ./cmd/ten-mien tu-dong --ma-cua-hang quochuy --xem-truoc   # chỉ xem, không ghi", False, True)], sau=6)
bang(
    ["Tên cửa hàng", "Địa chỉ được cấp", "Vì sao"],
    [
        ["Quốc Huy", "`quochuy.selliotech.store`", "bỏ dấu, bỏ khoảng trắng, viết liền"],
        ["Tạp hoá Dì Tư", "`taphoaditu.selliotech.store`", "cùng quy tắc"],
        ["Quoc Huy (khách thứ hai trùng tên)", "`quochuy-2.selliotech.store`", "nhãn đã có chủ thì thêm hậu tố"],
        ["Order", "`order-2.selliotech.store`", "`order` là nhãn nền tảng giữ lại"],
        ["“!!! ???” (không rút được chữ nào)", "`<mã cửa hàng>.selliotech.store`", "lùi về mã cửa hàng"],
    ],
    rong=[1.9, 2.4, 2.2],
)
cham([("Quy tắc sinh nhãn nằm ở `domain.NhanTenMienTuTen`, KHÔNG viết trong công cụ dòng lệnh: ngày có luồng đăng ký tự phục vụ thì nó gọi đúng hàm đó và ra đúng địa chỉ này. Hai chỗ tự bỏ dấu theo cách riêng là hai địa chỉ khác nhau cho cùng một khách.", False, False)])
cham([("Chạy lại lệnh KHÔNG đẻ thêm địa chỉ thứ hai — cửa hàng đã có subdomain thì lệnh dừng và nói ra. Muốn thêm thật thì `ten-mien them --host`.", False, False)])
cham([("Nhãn nền tảng (`order`, `admin`, `api`, `www`, `mail`, `cdn`…) được giữ lại: cấp mất rồi đòi lại nghĩa là đổi địa chỉ của một cửa hàng đang chạy — thứ khách đã in lên card và bao bì.", False, False)])
cham([("Nhãn cắt ở 30 ký tự (bằng độ dài `tenants.code`) — đủ để đọc qua điện thoại.", False, False)])

khung("Cấp ngoài luật thì phải nói lý do:",
      "`--dac-cach \"<lý do>\"` bỏ qua cả ba điều kiện, nhưng lý do được GHI VÀO cột `note` chứ không chỉ in ra màn hình. "
      "Sáu tháng sau nhìn một tên miền không khớp gói nào, người đọc phải trả lời được “ngoại lệ có chủ đích hay chỗ hở”.", mau=DO)
khung("Hệ quả ngay hôm nay:", "`subscriptions` đang rỗng nên MỌI lượt `ten-mien them` đều bị từ chối trừ khi dùng `--dac-cach`. "
      "Đó là kết quả đúng — sổ trống thì không cửa hàng nào chứng minh được mình đã mua gói gì.", mau=XAM)

# ---- 3.6 platform_users
h2("3.6  platform_users — tài khoản khu điều hành")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Người của NỀN TẢNG: mình và người làm cùng, đăng nhập vào `admin.selliotech.store` để mở/khoá cửa hàng, xem thuê bao."],
        ["Chức năng", "Nguồn xác thực cho SaaS Console qua `POST /auth/platform-login` (đăng nhập bằng email, không phải 3 ô)."],
        ["Cột đáng chú ý", "`email` · `full_name` · `password_hash` (bcrypt) · `role` = `owner` | `operator` | `support` · `status` = `active` | `locked`"],
        ["Máy cục bộ", "0 dòng"],
        ["Máy thật", "0 dòng"],
    ],
    rong=[1.3, 5.2],
)
khung("Cả hai máy đều 0 dòng:", "nghĩa là hiện KHÔNG ai đăng nhập được vào khu điều hành, và trong repo chưa có lệnh nào tạo tài khoản này "
      "(`cmd/tao-admin` chỉ tạo tài khoản cửa hàng bên data plane). Tạo tài khoản đầu tiên là việc phải làm trước khi khu điều hành dùng được trên máy thật.", mau=DO)
cham([("Không dùng chung bảng `users` với vai trò super_admin là chủ ý: tài khoản nằm chung bảng với người của MỘT cửa hàng thì nó mang `tenant_id` của cửa hàng đó, và khi plugin GORM chèn `WHERE tenant_id = ?` vào mọi truy vấn thì chính người điều hành bị nhốt trong tenant 1.", False, False)])
cham([("Bảng không có cột `tenant_id`, và đó là toàn bộ lý do nó tồn tại.", False, False)])

# ---- 3.7 schema_migrations
h2("3.7  schema_migrations — sổ migration của riêng lược đồ này")
bang(
    ["", ""],
    [
        ["Ý nghĩa", "Mỗi dòng là một tệp `.sql` đã chạy trên database này: số hiệu, tên, chạy lúc nào, mất bao lâu, và vân tay nội dung tệp."],
        ["Chức năng", "Cho công cụ `migrate` biết còn tệp nào chưa chạy, và phát hiện tệp đã chạy bị sửa nội dung sau đó."],
        ["Máy cục bộ", "4 dòng (`0001` control-plane, `0002` apps, `0003` plans, `0004` tên miền theo gói)"],
        ["Máy thật", "1 dòng (`0001` control-plane)"],
    ],
    rong=[1.3, 5.2],
)
khung("Đừng sửa tay bảng này, và đừng sửa tệp .sql đã chạy.",
      "Công cụ giữ vân tay; tệp đã chạy mà đổi nội dung sẽ báo lệch, và `deploy` gọi `migrate` để chặn triển khai khi lệch. "
      "Cần thêm gì thì viết tệp migration MỚI.", mau=DO)
p("Mỗi lược đồ có bảng `schema_migrations` RIÊNG, đánh số riêng từ `0001`. `database/migrations` là của data plane, "
  "`database/platform` là của control plane — chạy nhầm cờ `-nen-tang` là chạy tệp của lược đồ này vào database của lược đồ kia.")

doc.add_page_break()

# ---------------------------------------------------------------- 4
h1("4. selliotech — dữ liệu bán hàng (43 bảng)")

p("Đây là database phần mềm dùng hằng ngày. Ba điều đúng với gần như mọi bảng dưới đây:")
cham([("`tenant_id` ", False, True), ("có trên 37/38 bảng gốc (trừ `roles`) — đây là cột cắt dữ liệu theo từng khách hàng.", False, False)])
cham([("`shop_id` ", False, True), ("có thêm trên 11 bảng giao dịch — cắt tiếp theo từng chi nhánh.", False, False)])
cham([("Xoá mềm: nhiều bảng có `deleted_at` cộng cột `deleted_mark` để mã/email chỉ cần duy nhất giữa các dòng ĐANG SỐNG.", False, False)])

nhom = [
    ("4.1  Khách hàng, chi nhánh, người dùng", [
        ("tenants", "Mỗi dòng là MỘT khách hàng. Bảng cha của mọi `tenant_id`; `code` là mã cửa hàng gõ ở ô 1 màn hình đăng nhập.", "2 dòng ở cả hai máy"),
        ("shops", "Chi nhánh của một khách. Gói Chuỗi là nhiều dòng ở đây.", "2 dòng ở cả hai máy"),
        ("users", "Người đăng nhập: chủ shop, nhân viên, và cả khách mua hàng. `username` là ô 2 của đăng nhập 3 ô.", "cục bộ 2 · máy thật 4"),
        ("roles", "Bộ vai trò mặc định dùng chung (bảng DUY NHẤT không có `tenant_id`).", "4 dòng"),
        ("role_labels", "Tên hiển thị của vai trò theo TỪNG cửa hàng — shop này gọi “Thu ngân”, shop kia gọi “Bán hàng”.", "8 dòng"),
        ("user_addresses", "Sổ địa chỉ giao hàng của khách mua.", "0 dòng"),
        ("email_verifications", "Mã OTP gửi qua email lúc đăng ký / đặt lại mật khẩu.", "0 dòng"),
    ]),
    ("4.2  Danh mục sản phẩm", [
        ("categories", "Ngành hàng / danh mục, có cấp cha con.", "0 dòng"),
        ("brands", "Thương hiệu.", "0 dòng"),
        ("products", "Sản phẩm — phần mô tả chung, không giữ tồn kho.", "0 dòng"),
        ("product_variants", "Biến thể bán thật (màu, size, mã SKU, giá). Cột `stock_quantity` ở đây vẫn là tồn kho ĐANG DÙNG.", "0 dòng"),
        ("product_images", "Ảnh của sản phẩm / biến thể.", "0 dòng"),
        ("variant_stocks", "Tồn kho tách theo TỪNG chi nhánh — đã dựng sẵn nhưng code CHƯA đọc; phải nạp lại trước khi bỏ `stock_quantity`.", "0 dòng"),
    ]),
    ("4.3  Giỏ hàng", [
        ("carts", "Giỏ của một người mua (hoặc một phiên khách vãng lai).", "0 dòng"),
        ("cart_items", "Từng dòng hàng trong giỏ.", "0 dòng"),
    ]),
    ("4.4  Khuyến mãi", [
        ("vouchers", "Mã giảm giá: điều kiện, hạn dùng, số lần dùng.", "0 dòng"),
        ("voucher_usages", "Ai đã dùng mã nào, đơn nào — chỗ ép giới hạn số lần.", "0 dòng"),
        ("promotions", "Chương trình khuyến mãi theo thời gian.", "0 dòng"),
        ("promotion_targets", "Chương trình đó áp cho sản phẩm / danh mục nào.", "0 dòng"),
    ]),
    ("4.5  Đơn hàng", [
        ("orders", "Đơn hàng — bảng giao dịch trung tâm, có cả `tenant_id` lẫn `shop_id`.", "0 dòng"),
        ("order_items", "Từng món trong đơn, kèm giá tại thời điểm bán.", "0 dòng"),
        ("order_status_history", "Nhật ký đổi trạng thái đơn: ai đổi, lúc nào, từ gì sang gì.", "0 dòng"),
    ]),
    ("4.6  Thanh toán", [
        ("payments", "Lần thu tiền của một đơn: phương thức, số tiền, trạng thái.", "0 dòng"),
    ]),
    ("4.7  Kho", [
        ("inventory_transactions", "Nhật ký mọi lần tồn kho thay đổi và vì sao (bán, trả, nhập, kiểm kho).", "0 dòng"),
    ]),
    ("4.8  Trả hàng (khách trả cho shop)", [
        ("order_returns", "Phiếu trả hàng / hoàn tiền.", "0 dòng"),
        ("order_return_items", "Trả từng MÓN, không phải trả cả đơn.", "0 dòng"),
        ("order_return_history", "Nhật ký duyệt / từ chối phiếu trả.", "0 dòng"),
    ]),
    ("4.9  Nhập hàng từ nhà cung cấp", [
        ("suppliers", "Nhà cung cấp.", "0 dòng"),
        ("purchase_orders", "Phiếu đặt hàng nhập.", "0 dòng"),
        ("purchase_order_items", "Từng món trên phiếu nhập, kèm giá vốn.", "0 dòng"),
        ("purchase_order_history", "Nhật ký trạng thái phiếu nhập.", "0 dòng"),
        ("purchase_returns", "Phiếu trả hàng NGƯỢC về nhà cung cấp.", "0 dòng"),
        ("purchase_return_items", "Từng món trả về nhà cung cấp.", "0 dòng"),
        ("purchase_return_history", "Nhật ký phiếu trả nhà cung cấp.", "0 dòng"),
    ]),
    ("4.10  Tương tác của người mua", [
        ("product_reviews", "Đánh giá sản phẩm.", "0 dòng"),
        ("wishlists", "Sản phẩm khách lưu để mua sau.", "0 dòng"),
        ("notifications", "Thông báo trong phần mềm.", "0 dòng"),
    ]),
    ("4.11  Marketing, cấu hình, nhật ký", [
        ("banners", "Ảnh quảng cáo trên trang bán hàng.", "0 dòng"),
        ("newsletter_subscribers", "Email đăng ký nhận tin.", "0 dòng"),
        ("contact_requests", "Yêu cầu liên hệ gửi từ trang bán hàng.", "0 dòng"),
        ("settings", "Cấu hình của từng cửa hàng (tên shop, logo, thông tin in hoá đơn…).", "0 dòng"),
        ("activity_logs", "Nhật ký thao tác: ai làm gì, lúc nào — dùng để truy vết.", "0 dòng"),
        ("schema_migrations", "Sổ migration của data plane. Máy cục bộ và máy thật đều 4 dòng.", "4 dòng"),
    ]),
]

for tieu_de, ds in nhom:
    h2(tieu_de)
    bang(["Bảng", "Ý nghĩa / chức năng", "Số dòng"],
         [["`" + t + "`", y, s] for t, y, s in ds],
         rong=[1.5, 4.0, 1.0])

khung("Ghi chú về số dòng:", "cột này chụp lúc viết tài liệu (12/08/2026). Máy thật đã có 2 cửa hàng và 4 người dùng "
      "nhưng CHƯA có sản phẩm hay đơn hàng nào — hệ thống đã dựng xong nhưng chưa ai bán hàng thật trên đó.",
      mau=XAM)

doc.add_page_break()

# ---------------------------------------------------------------- 5
h1("5. Máy cục bộ và máy thật khác nhau chỗ nào (12/08/2026)")

bang(
    ["Thứ", "Máy cục bộ", "Máy thật", "Nghĩa là"],
    [
        ["Mã nguồn", "`5e3b946` + phần đang sửa chưa commit", "`5e3b946`", "Máy thật đang chạy đúng commit mới nhất đã đẩy lên."],
        ["`selliotech` — số bảng", "43", "43", "Data plane khớp nhau."],
        ["`selliotech` — migration", "4", "4", "Không nợ migration nào."],
        ["`selliotech_platform` — số bảng", "7", "5", "Máy thật THIẾU `apps` và `plans`."],
        ["`selliotech_platform` — migration", "4", "1", "Máy thật mới chạy `0001`; `0002`–`0004` còn nằm ở máy này."],
    ],
    rong=[1.9, 1.7, 1.1, 2.0],
)

p("Vì sao thiếu: ba tệp migration `0002_apps`, `0003_plans` và `0004_ten-mien-theo-goi` mới viết, chưa commit và chưa deploy. "
  "Script `deploy/scripts/02-trien-khai.sh` tự chạy migration cho CẢ HAI lược đồ, nên lần deploy tới chúng lên cùng — không phải làm gì thêm bằng tay.")

h2("Cách tự kiểm chứng hai máy có khớp nhau không")

khung("VÂN TAY LƯỢC ĐỒ KHÔNG SO ĐƯỢC GIỮA HAI MÁY NÀY.",
      "Máy cục bộ chạy MariaDB 10.4 (XAMPP), máy thật chạy MySQL 8.0. Hai engine mô tả CÙNG một lược đồ bằng chữ khác nhau, "
      "nên `van-tay` luôn ra hai chuỗi khác — và đó KHÔNG phải dấu hiệu database lệch.", mau=DO)

p("Ba khác biệt đã kiểm chứng ngày 12/08/2026, đều là cách viết chứ không phải lược đồ:")
bang(
    ["Thứ", "MariaDB 10.4 (cục bộ)", "MySQL 8.0 (máy thật)"],
    [
        ["Kiểu số nguyên", "`bigint(20) unsigned`", "`bigint unsigned` (MySQL 8 bỏ độ rộng hiển thị)"],
        ["Mặc định NULL", "chuỗi `NULL`", "NULL thật"],
        ["Mặc định chuỗi", "`'planned'` (có nháy)", "`planned` (không nháy)"],
    ],
    rong=[1.5, 2.3, 2.7],
)
p("`van-tay` / `so-sanh` vẫn đúng và vẫn nên dùng — nhưng chỉ giữa hai máy CÙNG engine "
  "(máy thử và máy thật, cả hai MySQL 8), hoặc để so cùng một máy trước và sau khi chạy migration.")

p("Giữa MariaDB và MySQL thì kiểm bằng hai câu này:", dam=True, truoc=6)
p("1. Số migration đã chạy — rẻ nhất, trả lời đúng câu “máy thật có thiếu migration nào không”:", sau=2)
rich([("SELECT MAX(version) FROM schema_migrations;   -- chạy ở cả hai máy, cả hai lược đồ", False, True)], sau=6)
p("2. So thẳng danh sách cột khi nghi ngờ có lệch thật:", sau=2)
rich([("SELECT table_name, column_name, is_nullable FROM information_schema.columns", False, True)], sau=0)
rich([("  WHERE table_schema = 'selliotech_platform' ORDER BY table_name, ordinal_position;", False, True)], sau=6)
p("Bỏ cột `column_type` ra khỏi câu so — chính nó là chỗ hai engine viết khác nhau. "
  "Thứ tự dòng cũng có thể lệch (`tenant_domains` và `tenants` sắp xếp khác nhau ở hai engine), nên sắp lại trước khi so.")

h2("Chạy công cụ Go trên máy thật")
p("`go run` bằng tài khoản root trên VPS sẽ đi tải lại toàn bộ thư viện và thường hỏng giữa chừng. "
  "Dùng đúng môi trường mà script triển khai dùng — cùng cache, cùng tài khoản:")
rich([("cd /var/www/selliotech/api && sudo -u selliotech env HOME=/var/www/selliotech \\", False, True)], sau=0)
rich([("  PATH=/usr/local/go/bin:$PATH GOCACHE=/var/cache/selliotech/go-build \\", False, True)], sau=0)
rich([("  GOMODCACHE=/var/cache/selliotech/go-mod /usr/local/go/bin/go run ./cmd/ten-mien danh-sach", False, True)], sau=6)

# ---------------------------------------------------------------- 6
h1("6. Quy tắc khi đụng vào database")

cham([("Lược đồ chỉ đổi bằng migration. ", True, False),
      ("Thêm/sửa bảng là viết tệp mới trong `database/migrations` (bán hàng) hoặc `database/platform` (nền tảng), rồi `go run ./cmd/migrate [-nen-tang] chay`.", False, False)])
cham([("Đừng chạy tệp .sql bằng DBeaver trên máy thật. ", True, False),
      ("Bảng sẽ lên nhưng `schema_migrations` không ghi, và lần deploy sau `migrate` báo lệch rồi chặn triển khai.", False, False)])
cham([("DBeaver dùng để ĐỌC và sửa DỮ LIỆU, không dùng để dựng lược đồ. ", True, False),
      ("Sửa giá trong `plans` bằng `UPDATE` thì hợp lệ — nhưng nhớ làm ở cả hai máy, vì đó là hai database riêng.", False, False)])
cham([("Đừng sửa tệp migration đã chạy. ", True, False),
      ("Vân tay sẽ lệch. Muốn đổi thì viết tệp mới.", False, False)])
cham([("Kiểm tra đang đứng ở database nào trước mỗi câu lệnh. ", True, False),
      ("`tenants` có ở cả hai database; gõ nhầm thì câu lệnh vẫn chạy, chỉ là sai bảng.", False, False)])
cham([("Máy thật có sao lưu tự động mỗi 12 giờ ", True, False),
      ("(`deploy/scripts/03-sao-luu.sh`), nhưng bản sao nằm cùng máy — nó cứu được xoá nhầm, không cứu được mất máy chủ.", False, False)])

# ---------------------------------------------------------------- 7
h1("7. Nhìn vào database thì thấy còn nợ những gì")

bang(
    ["Việc", "Bằng chứng nhìn thấy trong database"],
    [
        ["Khu điều hành trên máy thật chưa đăng nhập được",
         "`platform_users` = 0 dòng ở cả hai máy, mà SaaS Console đăng nhập qua `/auth/platform-login` đọc đúng bảng đó. Repo cũng chưa có lệnh nào tạo tài khoản này."],
        ["Hai cửa hàng đang chạy chưa vào sổ nền tảng",
         "`selliotech.tenants` có 2 dòng (`cuahang`, `quochuy`) nhưng `selliotech_platform.tenants` = 0 dòng. Hai bảng đang trôi khỏi nhau ngay từ bây giờ."],
        ["`cmd/tao-admin` chưa ghi sang control plane",
         "Hệ quả trực tiếp của dòng trên: tạo cửa hàng mới chỉ ghi một bên."],
        ["`subscriptions` chưa có `app_id`",
         "Bảng đang 0 dòng — làm bây giờ thì không phải điền ngược cho hợp đồng nào."],
        ["`variant_stocks` chưa được dùng",
         "0 dòng, trong khi tồn kho vẫn đọc/ghi `product_variants.stock_quantity`. Phải nạp lại bảng này trước khi bỏ cột kia."],
        ["Chưa cấp được tên miền nào cho tới khi có thuê bao",
         "Luật mới đòi thuê bao `active` + gói `own_domain = 1`, mà `subscriptions` = 0 dòng. Đường tạm thời là `--dac-cach` kèm lý do."],
        ["`apps` / `plans` / cột `own_domain` chưa lên máy thật",
         "Chỉ cần deploy, không cần thao tác tay."],
    ],
    rong=[2.2, 4.4],
)

# ---------------------------------------------------------------- 8
h1("8. Xoá một cửa hàng khỏi database")

p("Không có nút nào làm việc này, và cố ý như vậy. Xoá một cửa hàng là xoá dữ liệu kinh doanh "
  "của một khách hàng trả tiền — nó phải là một việc có người ngồi gõ tay, đọc lại con số trước "
  "khi bấm Enter.")

khung("ĐỌC TRƯỚC:",
      "gần như mọi trường hợp KHÔNG nên xoá. Cách đúng là `UPDATE tenants SET status = 'suspended'` "
      "— chặn đăng nhập, giữ nguyên dữ liệu, đóng tiền lại là mở. Chỉ xoá thật khi cửa hàng được "
      "tạo nhầm hoặc là dữ liệu thử.")

h2("Vì sao không xoá được bằng một câu DELETE")

cham([("Hai database, không có khoá ngoại bắc qua. ", True, False),
      ("Cửa hàng nằm ở `selliotech` (tenants · shops · users · toàn bộ hàng hoá, đơn từ) VÀ ở "
       "`selliotech_platform` (sổ khách hàng · hợp đồng · hoá đơn · tên miền). Xoá một bên thì "
       "bên kia không báo lỗi gì cả, nó chỉ lặng lẽ trỏ vào một id không còn ai.", False, False)])
cham([("`selliotech_platform.tenants.id` là số CHÉP từ data plane, ", True, False),
      ("không tự tăng. Cùng một cửa hàng mang cùng một id ở hai lược đồ — đó là dây nối duy nhất "
       "giữa hai bên.", False, False)])
cham([("42 khoá ngoại trỏ vào `tenants`, TẤT CẢ đều RESTRICT. ", True, False),
      ("Không cái nào CASCADE. Nghĩa là câu `DELETE FROM tenants` sẽ bị từ chối cho tới khi mọi "
       "dòng con biến mất — database giữ hộ, nhưng nó không dọn hộ.", False, False)])
cham([("40 bảng ở data plane có cột `tenant_id`. ", True, False),
      ("Viết tay danh sách đó vào tài liệu là bảo đảm ngày mai thêm bảng thứ 41 thì tài liệu sai. "
       "Vì vậy bước 4 dưới đây SINH danh sách từ `information_schema` chứ không chép sẵn.", False, False)])

h2("Quy trình — 5 bước")

p("Chạy theo đúng thứ tự này. Control plane TRƯỚC, data plane SAU: hợp đồng trỏ vào cửa hàng, "
  "nên bỏ cửa hàng trước là để lại hợp đồng mồ côi trong khoảng giữa hai lệnh.", sau=8)

p("Bước 1 — Sao lưu.", dam=True, sau=2)
p("Trên máy thật: `bash deploy/scripts/03-sao-luu.sh`. Trên máy cục bộ thì `mysqldump` hai lược đồ "
  "ra tệp. Không có bước này thì bốn bước sau là một chiều.", sau=8)

p("Bước 2 — Nhìn xem sắp xoá cái gì.", dam=True, sau=2)
rich([("SET @t := <tenant_id>;", False, True)], sau=0)
rich([("SELECT id, code, name, status FROM selliotech.tenants WHERE id = @t;", False, True)], sau=0)
rich([("SELECT s.id, s.plan, s.status,", False, True)], sau=0)
rich([("       (SELECT COUNT(*) FROM selliotech_platform.invoices i WHERE i.subscription_id = s.id) AS so_hoa_don", False, True)], sau=0)
rich([("  FROM selliotech_platform.subscriptions s WHERE s.tenant_id = @t;", False, True)], sau=6)
p("Có hoá đơn đã thu thì DỪNG LẠI. Đó là tiền đã vào sổ; xoá đi là báo cáo doanh thu của những "
  "tháng đã qua tự đổi số. Khoá cửa hàng, đừng xoá.", sau=8)

p("Bước 3 — Xoá ở control plane (`selliotech_platform`).", dam=True, sau=2)
p("Bốn bảng, con trước cha. `invoices` trỏ vào `subscriptions` nên phải đi đầu.", sau=2)
rich([("SET @t := <tenant_id>;", False, True)], sau=0)
rich([("START TRANSACTION;", False, True)], sau=0)
rich([("DELETE i FROM invoices i JOIN subscriptions s ON s.id = i.subscription_id WHERE s.tenant_id = @t;", False, True)], sau=0)
rich([("DELETE FROM tenant_domains WHERE tenant_id = @t;", False, True)], sau=0)
rich([("DELETE FROM subscriptions  WHERE tenant_id = @t;", False, True)], sau=0)
rich([("DELETE FROM tenants        WHERE id = @t;", False, True)], sau=0)
rich([("COMMIT;", False, True)], sau=6)
p("Câu `DELETE i FROM ... JOIN` đòi phải CÓ database đang chọn — chạy bằng `mysql -D selliotech_platform`, "
  "hoặc gõ `USE selliotech_platform;` trước. Thiếu thì nó báo “No database selected” dù tên bảng đã ghi đủ.", sau=8)

p("Bước 4 — Xoá ở data plane (`selliotech`).", dam=True, sau=2)
p("Sinh danh sách câu lệnh từ chính lược đồ, để không bảng nào bị bỏ sót:", sau=2)
rich([("SELECT GROUP_CONCAT(CONCAT('DELETE FROM `', TABLE_NAME, '` WHERE tenant_id = @t;') SEPARATOR '\\n')", False, True)], sau=0)
rich([("  FROM information_schema.COLUMNS", False, True)], sau=0)
rich([(" WHERE TABLE_SCHEMA = 'selliotech' AND COLUMN_NAME = 'tenant_id';", False, True)], sau=6)
p("Chép kết quả (40 câu) rồi bọc lại như sau. `FOREIGN_KEY_CHECKS = 0` là để khỏi phải sắp 40 bảng "
  "theo đúng thứ tự phụ thuộc — an toàn vì mọi dòng của cửa hàng này đều bị xoá trong CÙNG một "
  "giao dịch, không có trạng thái nửa vời nào lọt ra ngoài:", sau=2)
rich([("SET @t := <tenant_id>;", False, True)], sau=0)
rich([("START TRANSACTION;", False, True)], sau=0)
rich([("SET FOREIGN_KEY_CHECKS = 0;", False, True)], sau=0)
rich([("  ...40 câu DELETE vừa sinh...", False, True)], sau=0)
rich([("DELETE FROM tenants WHERE id = @t;", False, True)], sau=0)
rich([("SET FOREIGN_KEY_CHECKS = 1;", False, True)], sau=0)
rich([("COMMIT;", False, True)], sau=6)
khung("CẢNH BÁO:",
      "tắt kiểm khoá ngoại nghĩa là database thôi giữ hộ. Bỏ sót một bảng thì dòng thừa nằm lại "
      "mà KHÔNG câu lệnh nào báo lỗi — đó chính là lý do bước sinh danh sách ở trên phải đọc từ "
      "`information_schema` thay vì chép tay.", mau=DO)
p("Chạy bằng dòng lệnh thì thêm `--raw`: `mysql -N -B` mặc định đổi ký tự xuống dòng thành hai ký "
  "tự `\\n`, và cả 40 câu dính thành một dòng không chạy được.", truoc=4)

p("Bước 5 — Đối chiếu lại.", dam=True, sau=2)
p("Hai câu. Câu đầu bắt dòng còn sót, câu sau bắt bản ghi mồ côi giữa hai lược đồ:", sau=2)
rich([("-- (a) còn bảng nào sót dòng của cửa hàng đã xoá không — sinh câu giống bước 4,", False, True)], sau=0)
rich([("--     đổi DELETE FROM thành SELECT COUNT(*) FROM", False, True)], sau=0)
rich([("-- (b) sổ khách hàng trỏ vào cửa hàng không còn tồn tại:", False, True)], sau=0)
rich([("SELECT pt.id, pt.code FROM selliotech_platform.tenants pt", False, True)], sau=0)
rich([("  LEFT JOIN selliotech.tenants t ON t.id = pt.id WHERE t.id IS NULL;", False, True)], sau=6)
p("Cả hai phải trả về 0 dòng.", sau=8)

h2("Mấy chỗ dễ sai")

bang(
    ["Hiện tượng", "Nguyên nhân"],
    [
        ["Xoá xong vẫn báo “mã cửa hàng đã có người dùng”",
         "Mới xoá ở `selliotech_platform.tenants`. Khoá duy nhất `uq_tenants_code` nằm ở "
         "`selliotech.tenants` bên data plane, và màn hình kiểm trùng đọc đúng bảng đó."],
        ["`DELETE FROM tenants` báo lỗi khoá ngoại",
         "Còn dòng con. Cả 42 khoá đều RESTRICT, không cái nào tự dọn."],
        ["Xoá nhầm database",
         "`tenants` có mặt ở CẢ HAI lược đồ với cột gần giống nhau. Câu lệnh vẫn chạy, chỉ là sai bảng."],
        ["Cửa hàng biến mất nhưng vẫn đăng nhập được",
         "Mới xoá `tenants`, chưa xoá `users`. Đường đăng nhập 3 ô tra `users` theo (tenant_id, username)."],
        ["Xoá tenant 1 (`cuahang`)",
         "Đó là bản cài gốc do migration 0002 tạo. Migration đã chạy nên nó KHÔNG tự dựng lại — "
         "phải tạo tay bằng `go run ./cmd/tao-admin`."],
    ],
    rong=[2.4, 4.2],
)

p("Quy trình trên đã chạy thử trọn vẹn trên máy cục bộ ngày 13/08/2026: dựng một cửa hàng bằng "
  "màn hình “Thêm tài khoản dùng thử” (5 dòng ở 2 database), rồi xoá theo đúng 5 bước — bước 5 "
  "trả về 0 dòng ở cả hai câu.", mau=XAM, truoc=6)

p()
p("— Hết —", mau=XAM)

doc.core_properties.title = "Selliotech — Bản đồ database"
doc.core_properties.subject = "Ý nghĩa và chức năng từng bảng, máy cục bộ và máy thật"
doc.core_properties.comments = "Số liệu đọc từ MySQL ngày 12/08/2026"

doc.save(OUT)
print("Đã ghi:", OUT)
